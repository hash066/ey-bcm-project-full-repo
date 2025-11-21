"""
Document processing utilities for extracting text and controls from various file formats.
"""

from typing import Dict, List, Optional, Any
import logging
from datetime import datetime
from pathlib import Path
import re

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Processor for various document formats to extract text and compliance controls."""

    def __init__(self):
        """Initialize the document processor."""
        self.control_patterns = [
            # Pattern for numbered controls (1.1, 1.1.1, etc.)
            r'\b\d+\.\d+(?:\.\d+)*\s+[A-Z][^.!?]*[.!?]',
            # Pattern for bullet point controls
            r'\b•\s+[A-Z][^.!?]*[.!?]',
            r'\b-\s+[A-Z][^.!?]*[.!?]',
            # Pattern for "Control" or "Policy" statements
            r'\b(?:Control|Policy|Standard|Requirement)\s+\d+\s*[:.]?\s+[A-Z][^.!?]*[.!?]',
            # Pattern for SHALL/MUST statements
            r'\b(?:SHALL|MUST|SHOULD)\s+[A-Z][^.!?]*[.!?]',
            # Pattern for security requirements
            r'\b(?:Organizations?|Systems?|Users?)\s+(?:shall|must|should)\s+[A-Z][^.!?]*[.!?]',
        ]

    def process_file(self, file_path: str, file_type: Optional[str] = None) -> Dict[str, Any]:
        """
        Process a file and extract text and metadata.

        Args:
            file_path: Path to the file to process
            file_type: Type of file (word, pdf, excel). If None, will be inferred from extension.

        Returns:
            Dictionary containing extracted text and metadata

        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file type is not supported
            Exception: For other processing errors
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Infer file type if not provided
        if file_type is None:
            file_type = self._infer_file_type(file_path)

        logger.info(f"Processing {file_type} file: {file_path}")
        print(f"DEBUG: Processing {file_type} file: {file_path}")

        try:
            if file_type.lower() == "word":
                return self.process_word(str(file_path))
            elif file_type.lower() == "pdf":
                return self.process_pdf(str(file_path))
            elif file_type.lower() == "excel":
                return self.process_excel(str(file_path))
            else:
                raise ValueError(f"Unsupported file type: {file_type}")

        except Exception as e:
            logger.error(f"Error processing {file_type} file {file_path}: {str(e)}")
            print(f"DEBUG: Error processing {file_type} file {file_path}: {str(e)}")
            raise Exception(f"Failed to process {file_type} file: {str(e)}")

    def process_word(self, file_path: str) -> Dict[str, Any]:
        """
        Process a Word document using python-docx.

        Args:
            file_path: Path to the Word document

        Returns:
            Dictionary with extracted text and metadata
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required for Word document processing")

        try:
            doc = Document(file_path)
            text_content = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text.strip())

            full_text = '\n'.join(text_content)

            # Extract potential controls
            controls = self.extract_controls(full_text)

            return {
                "text": full_text,
                "controls": controls,
                "metadata": {
                    "type": "word",
                    "file_path": file_path,
                    "paragraphs": len(doc.paragraphs),
                    "tables": len(doc.tables),
                    "extracted_at": datetime.utcnow().isoformat(),
                    "controls_found": len(controls)
                }
            }

        except Exception as e:
            logger.error(f"Error processing Word document {file_path}: {str(e)}")
            raise Exception(f"Word processing failed: {str(e)}")

    def process_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        Process a PDF document using PyPDF2.

        Args:
            file_path: Path to the PDF document

        Returns:
            Dictionary with extracted text and metadata
        """
        try:
            import PyPDF2
        except ImportError:
            raise ImportError("PyPDF2 is required for PDF document processing")

        try:
            text_content = []

            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)

                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            text_content.append(f"--- Page {page_num} ---\n{text.strip()}")
                    except Exception as e:
                        logger.warning(f"Could not extract text from page {page_num}: {str(e)}")
                        text_content.append(f"--- Page {page_num} ---\n[Text extraction failed]")

            full_text = '\n\n'.join(text_content)

            # Extract potential controls
            controls = self.extract_controls(full_text)

            return {
                "text": full_text,
                "controls": controls,
                "metadata": {
                    "type": "pdf",
                    "file_path": file_path,
                    "pages": len(pdf_reader.pages),
                    "extracted_at": datetime.utcnow().isoformat(),
                    "controls_found": len(controls)
                }
            }

        except Exception as e:
            logger.error(f"Error processing PDF document {file_path}: {str(e)}")
            raise Exception(f"PDF processing failed: {str(e)}")

    def process_excel(self, file_path: str) -> Dict[str, Any]:
        """
        Process an Excel document using openpyxl.

        Args:
            file_path: Path to the Excel document

        Returns:
            Dictionary with extracted text and metadata
        """
        try:
            import openpyxl
        except ImportError:
            raise ImportError("openpyxl is required for Excel document processing")

        try:
            workbook = openpyxl.load_workbook(file_path, read_only=True)
            text_content = []
            sheet_info = []

            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_text = [f"--- Sheet: {sheet_name} ---"]

                for row_num, row in enumerate(sheet.iter_rows(values_only=True), 1):
                    row_text = []
                    for cell_value in row:
                        if cell_value is not None:
                            # Convert to string and clean up
                            cell_str = str(cell_value).strip()
                            if cell_str:
                                row_text.append(cell_str)

                    if row_text:
                        sheet_text.append(f"Row {row_num}: {' | '.join(row_text)}")

                if len(sheet_text) > 1:  # More than just the sheet header
                    text_content.extend(sheet_text)
                    sheet_info.append({
                        "name": sheet_name,
                        "rows": sheet.max_row,
                        "columns": sheet.max_column
                    })

            full_text = '\n'.join(text_content)

            # Extract potential controls
            controls = self.extract_controls(full_text)

            return {
                "text": full_text,
                "controls": controls,
                "metadata": {
                    "type": "excel",
                    "file_path": file_path,
                    "sheets": sheet_info,
                    "total_sheets": len(workbook.sheetnames),
                    "extracted_at": datetime.utcnow().isoformat(),
                    "controls_found": len(controls)
                }
            }

        except Exception as e:
            logger.error(f"Error processing Excel document {file_path}: {str(e)}")
            raise Exception(f"Excel processing failed: {str(e)}")

    def extract_controls(self, text: str) -> List[str]:
        """
        Extract potential compliance controls from text.

        Args:
            text: Text to analyze for control statements

        Returns:
            List of potential control statements
        """
        if not text:
            print("DEBUG: No text provided for control extraction")
            return []

        print(f"DEBUG: Extracting controls from {len(text)} characters of text")
        print(f"DEBUG: First 200 characters for control extraction: {text[:200]}...")

        controls = []
        lines = text.split('\n')

        for line in lines:
            line = line.strip()
            if not line or len(line) < 10:  # Skip very short lines
                continue

            # Check each pattern
            for pattern in self.control_patterns:
                matches = re.findall(pattern, line, re.MULTILINE)
                for match in matches:
                    # Clean up the match
                    clean_match = match.strip()
                    if (len(clean_match) > 20 and  # Must be substantial
                        clean_match not in controls and  # Avoid duplicates
                        not clean_match.isupper()):  # Avoid all caps lines (likely headers)
                        controls.append(clean_match)

        # Remove duplicates while preserving order
        seen = set()
        unique_controls = []
        for control in controls:
            if control not in seen:
                seen.add(control)
                unique_controls.append(control)

        print(f"DEBUG: Extracted {len(unique_controls)} potential controls from text")
        logger.info(f"Extracted {len(unique_controls)} potential controls from text")
        return unique_controls

    def _infer_file_type(self, file_path: Path) -> str:
        """
        Infer file type from file extension.

        Args:
            file_path: Path to the file

        Returns:
            File type string
        """
        extension = file_path.suffix.lower()
        logger.info(f"Detected file extension: '{extension}' for file: {file_path}")

        if extension in ['.docx', '.doc']:
            logger.info(f"Inferred file type 'word' for extension {extension}")
            return "word"
        elif extension == '.pdf':
            logger.info(f"Inferred file type 'pdf' for extension {extension}")
            return "pdf"
        elif extension in ['.xlsx', '.xls']:
            logger.info(f"Inferred file type 'excel' for extension {extension}")
            return "excel"
        else:
            logger.error(f"Cannot infer file type from extension: {extension}")
            raise ValueError(f"Cannot infer file type from extension: {extension}")

    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions."""
        return ['.docx', '.doc', '.pdf', '.xlsx', '.xls']

